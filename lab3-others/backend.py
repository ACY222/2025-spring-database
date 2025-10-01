from flask import Flask, request, jsonify, render_template, make_response
import mysql.connector
from mysql.connector import Error
import docx
import xlsxwriter
from io import BytesIO
import os

app = Flask(__name__)

# 数据库配置
app.config['MYSQL_HOST'] = 'laptop-AC'
app.config['MYSQL_USER'] = 'ac'
app.config['MYSQL_PASSWORD'] = 'yuanNBSBXL2+'
app.config['MYSQL_DB'] = 'lab3'

# 连接数据库
def get_db_connection():
    try:
        conn = mysql.connector.connect(
            host=app.config['MYSQL_HOST'],
            user=app.config['MYSQL_USER'],
            password=app.config['MYSQL_PASSWORD'],
            database=app.config['MYSQL_DB']
        )
        return conn
    except Error as e:
        print(f"数据库连接错误: {e}")
        return None

# 教师查询API
@app.route('/api/teachers', methods=['GET'])
def teachers():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    
    try:
        teacher_id = request.args.get('id')
        if teacher_id:
            # 查询单个教师信息
            cursor.execute("SELECT * FROM teacher WHERE teacher_id = %s", (teacher_id,))
            teacher = cursor.fetchone()
            if not teacher:
                return jsonify({"success": False, "error": "教师不存在"}), 404
            
            # 获取教师参与的论文数量
            cursor.execute(
                "SELECT COUNT(*) as paper_count FROM teacher_paper WHERE teacher_id = %s",
                (teacher_id,)
            )
            paper_count = cursor.fetchone()['paper_count']
            
            # 获取教师参与的项目数量
            cursor.execute(
                "SELECT COUNT(*) as project_count FROM teacher_project WHERE teacher_id = %s",
                (teacher_id,)
            )
            project_count = cursor.fetchone()['project_count']
            
            teacher['paper_count'] = paper_count
            teacher['project_count'] = project_count
            
            return jsonify({"success": True, "data": teacher})
        else:
            # 查询所有教师
            cursor.execute("SELECT * FROM teacher")
            teachers = cursor.fetchall()
            return jsonify({"success": True, "data": teachers})
    except Error as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        cursor.close()
        conn.close()

# 论文管理API
@app.route('/api/papers', methods=['GET', 'POST', 'PUT', 'DELETE'])
def papers():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    
    if request.method == 'GET':
        # 查询所有论文或按ID查询
        paper_id = request.args.get('id')
        if paper_id:
            cursor.execute("SELECT * FROM paper WHERE paper_id = %s", (paper_id,))
        else:
            cursor.execute("SELECT * FROM paper")
        papers = cursor.fetchall()
        return jsonify(papers)
    
    elif request.method == 'POST':
        # 添加论文
        data = request.json
        try:
            # 插入论文基本信息
            cursor.execute(
                "INSERT INTO paper (title, publish_year, level, source, type) "
                "VALUES (%s, %s, %s, %s, %s)",
                (data['title'], data['publish_year'], data['level'], 
                 data['source'], data['type'])
            )
            paper_id = cursor.lastrowid
            
            # 处理作者信息
            authors = data.get('authors', [])
            if authors:
                # 检查通讯作者数量
                corr_count = sum(1 for a in authors if a.get('is_corresponding', False))
                if corr_count != 1:
                    conn.rollback()
                    return jsonify({"error": "一篇论文必须有且仅有一位通讯作者"}), 400
                
                # 检查排名是否重复
                ranks = [a['rank'] for a in authors]
                if len(set(ranks)) != len(ranks):
                    conn.rollback()
                    return jsonify({"error": "论文作者排名不能重复"}), 400
                
                # 插入作者关系
                for author in authors:
                    cursor.execute(
                        "INSERT INTO teacher_paper (teacher_id, paper_id, rank, is_corresponding) "
                        "VALUES (%s, %s, %s, %s)",
                        (author['teacher_id'], paper_id, author['rank'], 
                         author.get('is_corresponding', False))
                    )
            
            conn.commit()
            return jsonify({"message": "论文添加成功", "paper_id": paper_id}), 201
        except Error as e:
            conn.rollback()
            return jsonify({"error": str(e)}), 500
        finally:
            cursor.close()
            conn.close()
    
    # PUT和DELETE方法类似，省略部分代码...

# 项目管理API
@app.route('/api/projects', methods=['GET', 'POST', 'PUT', 'DELETE'])
def projects():
    # 实现项目的增删改查，包含经费验证
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    
    if request.method == 'POST':
        data = request.json
        try:
            # 插入项目基本信息
            cursor.execute(
                "INSERT INTO project (project_name, source, project_type, "
                "total_fund, start_year, end_year) "
                "VALUES (%s, %s, %s, %s, %s, %s)",
                (data['project_name'], data['source'], data['project_type'],
                 data['total_fund'], data['start_year'], data['end_year'])
            )
            project_id = cursor.lastrowid
            
            # 处理承担项目的教师
            teachers = data.get('teachers', [])
            if teachers:
                total_fund = 0
                for teacher in teachers:
                    total_fund += teacher['fund']
                
                # 验证经费总和
                if total_fund != data['total_fund']:
                    conn.rollback()
                    return jsonify({"error": "承担经费总额必须等于项目总经费"}), 400
                
                # 检查排名是否重复
                ranks = [t['rank'] for t in teachers]
                if len(set(ranks)) != len(ranks):
                    conn.rollback()
                    return jsonify({"error": "项目成员排名不能重复"}), 400
                
                # 插入教师项目关系
                for teacher in teachers:
                    cursor.execute(
                        "INSERT INTO teacher_project (teacher_id, project_id, fund, rank) "
                        "VALUES (%s, %s, %s, %s)",
                        (teacher['teacher_id'], project_id, teacher['fund'], teacher['rank'])
                    )
            
            conn.commit()
            return jsonify({"message": "项目添加成功", "project_id": project_id}), 201
        except Error as e:
            conn.rollback()
            return jsonify({"error": str(e)}), 500
        finally:
            cursor.close()
            conn.close()
    
    # 其他方法省略...

# 课程管理API
@app.route('/api/courses', methods=['GET', 'POST', 'PUT', 'DELETE'])
def courses():
    # 实现课程的增删改查，包含学时验证
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    
    if request.method == 'POST':
        data = request.json
        try:
            # 插入课程基本信息
            cursor.execute(
                "INSERT INTO course (course_id, course_name, credit_hours, course_nature) "
                "VALUES (%s, %s, %s, %s)",
                (data['course_id'], data['course_name'], data['credit_hours'], data['course_nature'])
            )
            
            # 处理主讲教师
            teachers = data.get('teachers', [])
            if teachers:
                total_hours = 0
                for teacher in teachers:
                    total_hours += teacher['teaching_hours']
                
                # 验证学时总和
                if total_hours != data['credit_hours']:
                    conn.rollback()
                    return jsonify({"error": "主讲学时总额必须等于课程总学时"}), 400
                
                # 插入教师课程关系
                for teacher in teachers:
                    cursor.execute(
                        "INSERT INTO teacher_course (teacher_id, course_id, semester, year, teaching_hours) "
                        "VALUES (%s, %s, %s, %s, %s)",
                        (teacher['teacher_id'], data['course_id'], teacher['semester'],
                         teacher['year'], teacher['teaching_hours'])
                    )
            
            conn.commit()
            return jsonify({"message": "课程添加成功"}), 201
        except Error as e:
            conn.rollback()
            return jsonify({"error": str(e)}), 500
        finally:
            cursor.close()
            conn.close()
    
    # 其他方法省略...

# 查询统计功能
@app.route('/api/statistics', methods=['GET'])
def statistics():
    teacher_id = request.args.get('teacher_id')
    start_year = request.args.get('start_year')
    end_year = request.args.get('end_year')
    
    if not teacher_id or not start_year or not end_year:
        return jsonify({"error": "请提供教师工号和年份范围"}), 400
    
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    
    try:
        # 查询教师基本信息
        cursor.execute(
            "SELECT * FROM teacher WHERE teacher_id = %s", (teacher_id,)
        )
        teacher = cursor.fetchone()
        if not teacher:
            return jsonify({"error": "教师不存在"}), 404
        
        # 查询教学情况
        cursor.execute(
            "SELECT c.course_id, c.course_name, tc.teaching_hours, "
            "tc.semester, tc.year "
            "FROM teacher_course tc "
            "JOIN course c ON tc.course_id = c.course_id "
            "WHERE tc.teacher_id = %s AND tc.year BETWEEN %s AND %s",
            (teacher_id, start_year, end_year)
        )
        teaching = cursor.fetchall()
        
        # 查询论文情况
        cursor.execute(
            "SELECT p.title, p.source, p.publish_year, "
            "CASE p.level "
            "WHEN 1 THEN 'CCF-A' WHEN 2 THEN 'CCF-B' WHEN 3 THEN 'CCF-C' "
            "WHEN 4 THEN '中文CCF-A' WHEN 5 THEN '中文CCF-B' ELSE '无级别' END AS level, "
            "tp.rank, tp.is_corresponding "
            "FROM teacher_paper tp "
            "JOIN paper p ON tp.paper_id = p.paper_id "
            "WHERE tp.teacher_id = %s AND p.publish_year BETWEEN %s AND %s "
            "ORDER BY p.publish_year DESC",
            (teacher_id, start_year, end_year)
        )
        papers = cursor.fetchall()
        
        # 查询项目情况
        cursor.execute(
            "SELECT p.project_name, p.source, "
            "CASE p.project_type "
            "WHEN 1 THEN '国家级项目' WHEN 2 THEN '省部级项目' "
            "WHEN 3 THEN '市厅级项目' WHEN 4 THEN '企业合作项目' ELSE '其它类型项目' END AS type, "
            "p.start_year, p.end_year, p.total_fund, tp.fund "
            "FROM teacher_project tp "
            "JOIN project p ON tp.project_id = p.project_id "
            "WHERE tp.teacher_id = %s AND p.start_year BETWEEN %s AND %s "
            "ORDER BY p.start_year DESC",
            (teacher_id, start_year, end_year)
        )
        projects = cursor.fetchall()
        
        result = {
            "teacher": teacher,
            "teaching": teaching,
            "papers": papers,
            "projects": projects
        }
        return jsonify(result)
    
    except Error as e:
        return jsonify({"error": str(e)}), 500
    finally:
        cursor.close()
        conn.close()

# # 导出功能（选做）
# @app.route('/api/export', methods=['GET'])
# def export():
#     format = request.args.get('format', 'docx')
#     teacher_id = request.args.get('teacher_id')
#     start_year = request.args.get('start_year')
#     end_year = request.args.get('end_year')
    
#     if not teacher_id or not start_year or not end_year:
#         return jsonify({"error": "请提供教师工号和年份范围"}), 400
    
#     conn = get_db_connection()
#     cursor = conn.cursor(dictionary=True)
    
#     try:
#         # 查询统计数据
#         cursor.execute(
#             "SELECT * FROM teacher WHERE teacher_id = %s", (teacher_id,)
#         )
#         teacher = cursor.fetchone()
#         if not teacher:
#             return jsonify({"error": "教师不存在"}), 404
        
#         # 查询教学、论文、项目数据（代码同上，省略）
        
#         # 导出为Word
#         if format == 'docx':
#             doc = docx.Document()
#             doc.add_heading(f'教师教学科研工作统计({start_year}-{end_year})', 0)
            
#             # 教师基本信息
#             doc.add_heading('教师基本信息', level=1)
#             doc.add_paragraph(f'工号: {teacher["teacher_id"]}')
#             doc.add_paragraph(f'姓名: {teacher["name"]}')
#             doc.add_paragraph(f'性别: {"男" if teacher["gender"] == 1 else "女"}')
#             doc.add_paragraph(f'职称: {get_title_name(teacher["title"])}')
            
#             # 教学情况
#             if teaching:
#                 doc.add_heading('教学情况', level=1)
#                 for course in teaching:
#                     semester_name = get_semester_name(course['semester'])
#                     doc.add_paragraph(f'课程号: {course["course_id"]} '
#                                      f'课程名: {course["course_name"]} '
#                                      f'主讲学时: {course["teaching_hours"]} '
#                                      f'学期: {course["year"]}{semester_name}')
            
#             # 发表论文情况
#             if papers:
#                 doc.add_heading('发表论文情况', level=1)
#                 for i, paper in enumerate(papers, 1):
#                     corr = "通讯作者" if paper["is_corresponding"] else ""
#                     doc.add_paragraph(f'{i}. {paper["title"]}, {paper["source"]}, '
#                                      f'{paper["publish_year"]}, {paper["level"]}, '
#                                      f'排名第{paper["rank"]}, {corr}')
            
#             # 承担项目情况
#             if projects:
#                 doc.add_heading('承担项目情况', level=1)
#                 for i, project in enumerate(projects, 1):
#                     doc.add_paragraph(f'{i}. {project["project_name"]}, {project["source"]}, '
#                                      f'{project["type"]}, {project["start_year"]}-{project["end_year"]}, '
#                                      f'总经费: {project["total_fund"]}, 承担经费: {project["fund"]}')
            
#             # 保存到内存并返回
#             buffer = BytesIO()
#             doc.save(buffer)
#             buffer.seek(0)
#             response = make_response(buffer.read())
#             response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#             response.headers['Content-Disposition'] = f'attachment; filename=teacher_statistics_{teacher_id}_{start_year}-{end_year}.docx'
#             return response
        
#         # 导出为Excel
#         elif format == 'xlsx':
#             buffer = BytesIO()
#             with xlsxwriter.Workbook(buffer) as workbook:
#                 # 实现Excel导出逻辑...（类似Word）
#             buffer.seek(0)
#             response = make_response(buffer.read())
#             response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
#             response.headers['Content-Disposition'] = f'attachment; filename=teacher_statistics_{teacher_id}_{start_year}-{end_year}.xlsx'
#             return response
        
#         else:
#             return jsonify({"error": "不支持的导出格式"}), 400
    
#     except Error as e:
#         return jsonify({"error": str(e)}), 500
#     finally:
#         cursor.close()
#         conn.close()

# 辅助函数：获取职称名称
def get_title_name(title_code):
    titles = {
        1: "博士后", 2: "助教", 3: "讲师", 4: "副教授", 5: "特任教授",
        6: "教授", 7: "助理研究员", 8: "特任副研究员", 9: "副研究员",
        10: "特任研究员", 11: "研究员"
    }
    return titles.get(title_code, f"未知职称({title_code})")

# 辅助函数：获取学期名称
def get_semester_name(semester_code):
    semesters = {1: "春", 2: "夏", 3: "秋"}
    return semesters.get(semester_code, f"未知学期({semester_code})")

# 前端页面路由
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/papers')
def papers_page():
    return render_template('papers.html')

# 其他页面路由...

if __name__ == '__main__':
    app.run(debug=True)
