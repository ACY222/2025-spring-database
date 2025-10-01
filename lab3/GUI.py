# GUI
import tkinter as tk
from tkinter import ttk, messagebox
from DAL import *
from tkinter import filedialog
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

current_teacher_id = None

## 发表论文情况
def paper():
    win = tk.Toplevel()
    win.title("发表论文情况")
    win.geometry("800x600")
    ttk.Label(win, text="请选择要进入的模块：", font=("Arial", 14)).pack(pady=20)

    ttk.Button(win, text="论文添加", width=20, command=add_paper).pack(pady=5)
    ttk.Button(win, text="论文删除", width=20, command=del_paper).pack(pady=5)
    ttk.Button(win, text="论文修改", width=20, command=renewpaper).pack(pady=5)
    ttk.Button(win, text="论文查询", width=20, command=searchpaper).pack(pady=5)

    win.mainloop()

def add_paper():
    sub_win = tk.Toplevel()
    sub_win.title("论文添加")
    sub_win.geometry("500x500")
    ttk.Label(sub_win, text=f"请输入指定信息：", font=("Arial", 16)).grid(row=0, column=0, columnspan=2, pady=10)

    ttk.Label(sub_win, text="论文序号：").grid(row=1, column=0, padx=10, pady=10, sticky="e")
    paper_id = ttk.Entry(sub_win)
    paper_id.grid(row=1, column=1)

    ttk.Label(sub_win, text="排名：").grid(row=2, column=0, padx=10, pady=10, sticky="e")
    rank = ttk.Entry(sub_win)
    rank.grid(row=2, column=1)

    ttk.Label(sub_win, text="是否通讯作者：").grid(row=3, column=0, padx=10, pady=10, sticky="e")
/2024_2Summer
/2024_3Fall
/2025_1Spring
/LaTeX_learn
/Linux
/Markdown
    is_cor = tk.StringVar()
    is_cor = ttk.Combobox(sub_win, values=[0, 1], state="readonly")
    is_cor.grid(row=3, column=1)
    is_cor.current(0)

    def submit():
        p_id = paper_id.get().strip()
        r = int(rank.get().strip())
        cor = int(is_cor.get().strip())
        if not paper_exists(p_id):
            insertpaper()
        if cor == 1 and one_cor_author(p_id):
            messagebox.showerror("错误", "该论文已经有一位通讯作者。")
            return
        if no_same_rank(p_id, r):
            messagebox.showerror("错误", "该排名已被占用。")
            return
        insert_teacher_post_paper(current_teacher_id, p_id, r, cor)
        messagebox.showinfo("成功", "论文作者信息添加成功！")

    submit_btn = ttk.Button(sub_win, text="提交", command=submit)
    submit_btn.grid(row=4, column=0, columnspan=2, pady=20)

    sub_win.mainloop()

def insertpaper():
    sub_win = tk.Toplevel()
    sub_win.title("请先登记论文")
    sub_win.geometry("300x500")

    ttk.Label(sub_win, text="论文序号：").grid(row=0, column=0, padx=10, pady=10, sticky="e")
    paper_id = ttk.Entry(sub_win)
    paper_id.grid(row=0, column=1)

    ttk.Label(sub_win, text="论文名称：").grid(row=1, column=0, padx=10, pady=10, sticky="e")
    paper_name = ttk.Entry(sub_win)
    paper_name.grid(row=1, column=1)

    ttk.Label(sub_win, text="发表源：").grid(row=2, column=0, padx=10, pady=10, sticky="e")
    source = ttk.Entry(sub_win)
    source.grid(row=2, column=1)

    ttk.Label(sub_win, text="发表年份：").grid(row=3, column=0, padx=10, pady=10, sticky="e")
    year = ttk.Entry(sub_win)
    year.grid(row=3, column=1)

    ttk.Label(sub_win, text="类型：").grid(row=4, column=0, padx=10, pady=10, sticky="e")
    type = tk.StringVar()
    type = ttk.Combobox(sub_win, values=[1, 2, 3, 4], state="readonly")
    type.grid(row=4, column=1)
    type.current(0)

    ttk.Label(sub_win, text="级别：").grid(row=5, column=0, padx=10, pady=10, sticky="e")
    level = tk.StringVar()
    level = ttk.Combobox(sub_win, values=[1, 2, 3, 4, 5, 6], state="readonly")
    level.grid(row=5, column=1)
    level.current(0)

    def insert():
        try:
            p_id = paper_id.get().strip()
            p_name = paper_name.get().strip()
            s = source.get().strip()
            y = year.get().strip()
            t = type.get()
            l = level.get()
            insert_paper(p_id, p_name, s, y, t, l)
            messagebox.showinfo("成功", "插入论文成功！")
        except Exception as e:
            messagebox.showerror("错误", f"插入失败：{e}")
            sub_win.destroy()

    submit_btn = ttk.Button(sub_win, text="提交", command=insert)
    submit_btn.grid(row=6, column=0, columnspan=2, pady=20)

    sub_win.mainloop()

def del_paper():
    sub_win = tk.Toplevel()
    sub_win.title("论文删除")
    sub_win.geometry("300x150")

    ttk.Label(sub_win, text="论文序号：").grid(row=0, column=0, padx=10, pady=10, sticky="e")
    paper_id = ttk.Entry(sub_win)
    paper_id.grid(row=0, column=1)

    def delete():
        try:
            p_id = int(paper_id.get())
            delete_teacher_post_paper(current_teacher_id, p_id)
            messagebox.showinfo("成功", "删除成功！")
        except Exception as e:
            messagebox.showerror("错误", f"删除失败：{e}")
            sub_win.destroy()

    submit_btn = ttk.Button(sub_win, text="提交", command=delete)
    submit_btn.grid(row=1, column=0, columnspan=2, pady=20)

    sub_win.mainloop()

def renewpaper():
    sub_win = tk.Toplevel()
    sub_win.title("论文修改")
    sub_win.geometry("300x500")
    ttk.Label(sub_win, text=f"请输入指定信息：", font=("Arial", 16))

    ttk.Label(sub_win, text="论文序号：").grid(row=0, column=0, padx=10, pady=10, sticky="e")
    paper_id = ttk.Entry(sub_win)
    paper_id.grid(row=0, column=1)

    ttk.Label(sub_win, text="论文名称：").grid(row=1, column=0, padx=10, pady=10, sticky="e")
    paper_name = ttk.Entry(sub_win)
    paper_name.grid(row=1, column=1)

    ttk.Label(sub_win, text="发表源：").grid(row=2, column=0, padx=10, pady=10, sticky="e")
    source = ttk.Entry(sub_win)
    source.grid(row=2, column=1)

    ttk.Label(sub_win, text="发表年份：").grid(row=3, column=0, padx=10, pady=10, sticky="e")
    year = ttk.Entry(sub_win)
    year.grid(row=3, column=1)

    ttk.Label(sub_win, text="类型：").grid(row=4, column=0, padx=10, pady=10, sticky="e")
    type = tk.StringVar()
    type = ttk.Combobox(sub_win, values=[1, 2, 3, 4], state="readonly")
    type.grid(row=4, column=1)
    type.current(0)

    ttk.Label(sub_win, text="级别：").grid(row=5, column=0, padx=10, pady=10, sticky="e")
    level = tk.StringVar()
    level = ttk.Combobox(sub_win, values=[1, 2, 3, 4, 5, 6], state="readonly")
    level.grid(row=5, column=1)
    level.current(0)

    def renew():
        try:
            p_id = paper_id.get().strip()
            p_name = paper_name.get().strip()
            s = source.get().strip()
            y = year.get().strip()
            t = type.get()
            l = level.get()
            renew_paper(p_id, p_name, s, y, t, l)
            messagebox.showinfo("成功", "更新论文成功！")
        except Exception as e:
            messagebox.showerror("错误", f"更新失败：{e}")
            sub_win.destroy()

    submit_btn = ttk.Button(sub_win, text="提交", command=renew)
    submit_btn.grid(row=6, column=0, columnspan=2, pady=20)

    sub_win.mainloop()

def searchpaper():
    sub_win = tk.Toplevel()
    sub_win.title("论文查询")
    sub_win.geometry("700x700")
    ttk.Label(sub_win, text=f"请输入查询条件（可为空）：", font=("Arial", 16))

    ttk.Label(sub_win, text="论文序号：").grid(row=0, column=0, padx=10, pady=10, sticky="e")
    paper_id = ttk.Entry(sub_win)
    paper_id.grid(row=0, column=1)

    ttk.Label(sub_win, text="论文名称：").grid(row=1, column=0, padx=10, pady=10, sticky="e")
    paper_name = ttk.Entry(sub_win)
    paper_name.grid(row=1, column=1)

    ttk.Label(sub_win, text="发表源：").grid(row=2, column=0, padx=10, pady=10, sticky="e")
    source = ttk.Entry(sub_win)
    source.grid(row=2, column=1)

    ttk.Label(sub_win, text="发表年份：").grid(row=3, column=0, padx=10, pady=10, sticky="e")
    year = ttk.Entry(sub_win)
    year.grid(row=3, column=1)

    ttk.Label(sub_win, text="类型：").grid(row=4, column=0, padx=10, pady=10, sticky="e")
    type = tk.StringVar()
    type = ttk.Combobox(sub_win, values=[1, 2, 3, 4], state="readonly")
    type.grid(row=4, column=1)
    type.set('') 

    ttk.Label(sub_win, text="级别：").grid(row=5, column=0, padx=10, pady=10, sticky="e")
    level = tk.StringVar()
    level = ttk.Combobox(sub_win, values=[1, 2, 3, 4, 5, 6], state="readonly")
    level.grid(row=5, column=1)
    level.set('')

    ttk.Label(sub_win, text="发表教师工号：").grid(row=6, column=0, padx=10, pady=10, sticky="e")
    teacher_id = ttk.Entry(sub_win)
    teacher_id.grid(row=6, column=1)

    columns = ["paper_id", "paper_name", "source", "post_year", "paper_type", "level", "teacher_id"]
    tree = ttk.Treeview(sub_win, columns=columns, show="headings", height=10)
    tree.grid(row=8, column=0, columnspan=2, pady=10)

    headings = ["论文编号", "论文名称", "发表源", "发表年份", "类型", "级别", "教师工号"]
    for col, heading in zip(columns, headings):
        tree.heading(col, text=heading)
        tree.column(col, width=100)

    def search():
        try:
            p_id = paper_id.get().strip()
            p_name = paper_name.get().strip()
            s = source.get().strip()
            y = year.get().strip()
            t = type.get()
            l = level.get()
            t_id = teacher_id.get()
            results = search_paper(p_id, p_name, s, y, t, l, t_id)

            for row in tree.get_children():
                tree.delete(row)
            for row in results:
                values = [row["paper_id"], row["paper_name"], row["source"],
                        row["post_year"], row["paper_type"], row["level"], 
                        row["teacher_id"]]
                tree.insert("", tk.END, values=values)
                # tree.insert("", tk.END, values=row)

        except Exception as e:
            messagebox.showerror("错误", f"查询失败：{e}")
            sub_win.destroy()

    submit_btn = ttk.Button(sub_win, text="提交", command=search)
    submit_btn.grid(row=7, column=0, columnspan=2, pady=20)

    sub_win.mainloop()

## 承担项目情况
def project():
    win = tk.Toplevel()
    win.title("承担项目情况")
    win.geometry("800x600")
    ttk.Label(win, text="请选择要进入的模块：", font=("Arial", 14)).pack(pady=20)

    ttk.Button(win, text="项目添加", width=20, command=add_project).pack(pady=5)
    ttk.Button(win, text="项目删除", width=20, command=del_project).pack(pady=5)
    ttk.Button(win, text="项目修改", width=20, command=renewproject).pack(pady=5)
    ttk.Button(win, text="项目查询", width=20, command=searchproject).pack(pady=5)

    win.mainloop()

def add_project():
    sub_win = tk.Toplevel()
    sub_win.title("项目添加")
    sub_win.geometry("500x500")
    ttk.Label(sub_win, text=f"请输入指定信息：", font=("Arial", 16)).grid(row=0, column=0, columnspan=2, pady=10)

    ttk.Label(sub_win, text="项目号：").grid(row=1, column=0, padx=10, pady=10, sticky="e")
    project_id = ttk.Entry(sub_win)
    project_id.grid(row=1, column=1)

    ttk.Label(sub_win, text="排名：").grid(row=2, column=0, padx=10, pady=10, sticky="e")
    rank = ttk.Entry(sub_win)
    rank.grid(row=2, column=1)

    ttk.Label(sub_win, text="承担经费：").grid(row=3, column=0, padx=10, pady=10, sticky="e")
    money = ttk.Entry(sub_win)
    money.grid(row=3, column=1)

    def submit():
        p_id = project_id.get().strip()
        r = int(rank.get().strip())
        m = int(money.get().strip())
        if not project_exists(p_id):
            insertproject()
        if no_same_project_rank(p_id, r):
            messagebox.showerror("错误", "该排名已被占用。")
            sub_win.destroy()
            return
        insert_teacher_hold_project(current_teacher_id, p_id, r, m)
        messagebox.showinfo("成功", "项目承担信息添加成功")
        if not money_correct(p_id):
            messagebox.showinfo("提示", "项目总金额已更新")

    submit_btn = ttk.Button(sub_win, text="提交", command=submit)
    submit_btn.grid(row=4, column=0, columnspan=2, pady=20)

    sub_win.mainloop()

def insertproject():
    sub_win = tk.Toplevel()
    sub_win.title("请先登记项目")
    sub_win.geometry("300x500")

    ttk.Label(sub_win, text="项目号：").grid(row=0, column=0, padx=10, pady=10, sticky="e")
    project_id = ttk.Entry(sub_win)
    project_id.grid(row=0, column=1)

    ttk.Label(sub_win, text="项目名称：").grid(row=1, column=0, padx=10, pady=10, sticky="e")
    project_name = ttk.Entry(sub_win)
    project_name.grid(row=1, column=1)

    ttk.Label(sub_win, text="项目来源：").grid(row=2, column=0, padx=10, pady=10, sticky="e")
    source = ttk.Entry(sub_win)
    source.grid(row=2, column=1)

    ttk.Label(sub_win, text="项目类型：").grid(row=3, column=0, padx=10, pady=10, sticky="e")
    type = tk.StringVar()
    type = ttk.Combobox(sub_win, values=[1, 2, 3, 4, 5], state="readonly")
    type.grid(row=3, column=1)
    type.current(0)

    ttk.Label(sub_win, text="开始年份：").grid(row=4, column=0, padx=10, pady=10, sticky="e")
    start_year = ttk.Entry(sub_win)
    start_year.grid(row=4, column=1)

    ttk.Label(sub_win, text="结束年份：").grid(row=5, column=0, padx=10, pady=10, sticky="e")
    end_year = ttk.Entry(sub_win)
    end_year.grid(row=5, column=1)

    def insert():
        try:
            p_id = project_id.get().strip()
            p_name = project_name.get().strip()
            s = source.get().strip()
            t = type.get()
            sy = start_year.get()
            ey = end_year.get()
            insert_project(p_id, p_name, s, t, 0, sy, ey)
            messagebox.showinfo("成功", "插入项目成功！")
        except Exception as e:
            messagebox.showerror("错误", f"插入失败：{e}")
            sub_win.destroy()

    submit_btn = ttk.Button(sub_win, text="提交", command=insert)
    submit_btn.grid(row=6, column=0, columnspan=2, pady=20)

    sub_win.mainloop()

def del_project():
    sub_win = tk.Toplevel()
    sub_win.title("项目删除")
    sub_win.geometry("300x150")

    ttk.Label(sub_win, text="项目号：").grid(row=0, column=0, padx=10, pady=10, sticky="e")
    project_id = ttk.Entry(sub_win)
    project_id.grid(row=0, column=1)

    def delete():
        try:
            p_id = project_id.get().strip()
            delete_teacher_hold_project(current_teacher_id, p_id)
            messagebox.showinfo("成功", "删除成功！")
            if not money_correct(p_id):
                messagebox.showinfo("提示", "项目总金额已更新")
        except Exception as e:
            messagebox.showerror("错误", f"删除失败：{e}")
            sub_win.destroy()

    submit_btn = ttk.Button(sub_win, text="提交", command=delete)
    submit_btn.grid(row=1, column=0, columnspan=2, pady=20)

    sub_win.mainloop()

def renewproject():
    sub_win = tk.Toplevel()
    sub_win.title("项目修改")
    sub_win.geometry("300x500")
    ttk.Label(sub_win, text=f"请输入指定信息：", font=("Arial", 16))

    ttk.Label(sub_win, text="项目号：").grid(row=0, column=0, padx=10, pady=10, sticky="e")
    project_id = ttk.Entry(sub_win)
    project_id.grid(row=0, column=1)

    ttk.Label(sub_win, text="项目名称：").grid(row=1, column=0, padx=10, pady=10, sticky="e")
    project_name = ttk.Entry(sub_win)
    project_name.grid(row=1, column=1)

    ttk.Label(sub_win, text="项目来源：").grid(row=2, column=0, padx=10, pady=10, sticky="e")
    source = ttk.Entry(sub_win)
    source.grid(row=2, column=1)

    ttk.Label(sub_win, text="项目类型：").grid(row=3, column=0, padx=10, pady=10, sticky="e")
    type = tk.StringVar()
    type = ttk.Combobox(sub_win, values=[1, 2, 3, 4, 5], state="readonly")
    type.grid(row=3, column=1)
    type.current(0)

    ttk.Label(sub_win, text="开始年份：").grid(row=5, column=0, padx=10, pady=10, sticky="e")
    start_year = ttk.Entry(sub_win)
    start_year.grid(row=5, column=1)

    ttk.Label(sub_win, text="结束年份：").grid(row=6, column=0, padx=10, pady=10, sticky="e")
    end_year = ttk.Entry(sub_win)
    end_year.grid(row=6, column=1)

    def renew():
        try:
            p_id = project_id.get().strip()
            p_name = project_name.get().strip()
            s = source.get().strip()
            t = type.get()
            sy = start_year.get()
            ey = end_year.get()
            renew_project(p_id, p_name, s, t, sy, ey)
            messagebox.showinfo("成功", "修改项目成功！")
            if not money_correct(p_id):
                messagebox.showinfo("提示", "项目总金额已更新")
        except Exception as e:
            messagebox.showerror("错误", f"修改失败：{e}")
            sub_win.destroy()

    submit_btn = ttk.Button(sub_win, text="提交", command=renew)
    submit_btn.grid(row=7, column=0, columnspan=2, pady=20)

    sub_win.mainloop()

def searchproject():
    sub_win = tk.Toplevel()
    sub_win.title("项目查询")
    sub_win.geometry("700x700")
    ttk.Label(sub_win, text=f"请输入查询条件（可为空）：", font=("Arial", 16))

    ttk.Label(sub_win, text="项目号：").grid(row=0, column=0, padx=10, pady=10, sticky="e")
    project_id = ttk.Entry(sub_win)
    project_id.grid(row=0, column=1)

    ttk.Label(sub_win, text="项目名称：").grid(row=1, column=0, padx=10, pady=10, sticky="e")
    project_name = ttk.Entry(sub_win)
    project_name.grid(row=1, column=1)

    ttk.Label(sub_win, text="项目来源：").grid(row=2, column=0, padx=10, pady=10, sticky="e")
    source = ttk.Entry(sub_win)
    source.grid(row=2, column=1)

    ttk.Label(sub_win, text="项目类型：").grid(row=3, column=0, padx=10, pady=10, sticky="e")
    type = tk.StringVar()
    type = ttk.Combobox(sub_win, values=[1, 2, 3, 4, 5], state="readonly")
    type.grid(row=3, column=1)
    type.set('') 

    ttk.Label(sub_win, text="总经费：").grid(row=4, column=0, padx=10, pady=10, sticky="e")
    money = ttk.Entry(sub_win)
    money.grid(row=4, column=1)

    ttk.Label(sub_win, text="开始年份：").grid(row=5, column=0, padx=10, pady=10, sticky="e")
    start_year = ttk.Entry(sub_win)
    start_year.grid(row=5, column=1)

    ttk.Label(sub_win, text="结束年份：").grid(row=6, column=0, padx=10, pady=10, sticky="e")
    end_year = ttk.Entry(sub_win)
    end_year.grid(row=6, column=1)

    ttk.Label(sub_win, text="承担教师工号：").grid(row=7, column=0, padx=10, pady=10, sticky="e")
    teacher_id = ttk.Entry(sub_win)
    teacher_id.grid(row=7, column=1)

    columns = ["project_id", "project_name", "project_source", "project_type", "project_money", "start_year", "end_year", "teacher_id"]
    tree = ttk.Treeview(sub_win, columns=columns, show="headings", height=10)
    tree.grid(row=9, column=0, columnspan=2, pady=10)

    headings = ["项目号", "项目名称", "项目来源", "项目类型", "总经费", "开始年份", "结束年份", "教师工号"]
    for col, heading in zip(columns, headings):
        tree.heading(col, text=heading)
        tree.column(col, width=100)

    def search():
        try:
            p_id = project_id.get().strip()
            p_name = project_name.get().strip()
            s = source.get().strip()
            t = type.get()
            m = money.get()
            sy = start_year.get()
            ey = end_year.get()
            t_id = teacher_id.get()
            results = search_project(p_id, p_name, s, t, m, sy, ey, t_id)

            for row in tree.get_children():
                tree.delete(row)
            for row in results:
                values = [row["project_id"], row["project_name"], row["project_source"],
                        row["project_type"], row["start_year"], row["end_year"], 
                        row["teacher_id"]]
                tree.insert("", tk.END, values=values)
                # tree.insert("", tk.END, values=row)

        except Exception as e:
            messagebox.showerror("错误", f"查询失败：{e}")
            sub_win.destroy()

    submit_btn = ttk.Button(sub_win, text="提交", command=search)
    submit_btn.grid(row=8, column=0, columnspan=2, pady=20)

    sub_win.mainloop()

## 主讲课程情况
def course():
    win = tk.Toplevel()
    win.title("主讲课程情况")
    win.geometry("800x600")
    ttk.Label(win, text="请选择要进入的模块：", font=("Arial", 14)).pack(pady=20)

    ttk.Button(win, text="课程添加", width=20, command=add_course).pack(pady=5)
    ttk.Button(win, text="课程删除", width=20, command=del_course).pack(pady=5)
    ttk.Button(win, text="课程修改", width=20, command=renewcourse).pack(pady=5)
    ttk.Button(win, text="课程查询", width=20, command=searchcourse).pack(pady=5)

    win.mainloop()

def add_course():
    sub_win = tk.Toplevel()
    sub_win.title("课程添加")
    sub_win.geometry("500x500")
    ttk.Label(sub_win, text=f"请输入指定信息：", font=("Arial", 16)).grid(row=0, column=0, columnspan=2, pady=10)

    ttk.Label(sub_win, text="课程号：").grid(row=1, column=0, padx=10, pady=10, sticky="e")
    course_id = ttk.Entry(sub_win)
    course_id.grid(row=1, column=1)

    ttk.Label(sub_win, text="年份：").grid(row=2, column=0, padx=10, pady=10, sticky="e")
    year = ttk.Entry(sub_win)
    year.grid(row=2, column=1)

    ttk.Label(sub_win, text="学期：").grid(row=3, column=0, padx=10, pady=10, sticky="e")
    term = tk.StringVar()
    term = ttk.Combobox(sub_win, values=[1, 2, 3], state="readonly")
    term.grid(row=3, column=1)
    term.current(0)

    ttk.Label(sub_win, text="承担学时：").grid(row=4, column=0, padx=10, pady=10, sticky="e")
    hour = ttk.Entry(sub_win)
    hour.grid(row=4, column=1)

    def submit():
        c_id = course_id.get().strip()
        y = int(year.get().strip())
        t = int(term.get().strip())
        h = int(hour.get().strip())
        if not course_exists(c_id):
            insertcourse()
        insert_teacher_teach_course(current_teacher_id, c_id, y, t, h)
        messagebox.showinfo("成功", "主讲课程信息添加成功！")
        if not hour_correct(c_id):
            messagebox.showinfo("提示", "课程总学时已更新")

    submit_btn = ttk.Button(sub_win, text="提交", command=submit)
    submit_btn.grid(row=5, column=0, columnspan=2, pady=20)

    sub_win.mainloop()

def insertcourse():
    sub_win = tk.Toplevel()
    sub_win.title("请先登记课程")
    sub_win.geometry("300x500")

    ttk.Label(sub_win, text="课程号：").grid(row=0, column=0, padx=10, pady=10, sticky="e")
    course_id = ttk.Entry(sub_win)
    course_id.grid(row=0, column=1)

    ttk.Label(sub_win, text="课程名称：").grid(row=1, column=0, padx=10, pady=10, sticky="e")
    course_name = ttk.Entry(sub_win)
    course_name.grid(row=1, column=1)

    ttk.Label(sub_win, text="课程性质：").grid(row=3, column=0, padx=10, pady=10, sticky="e")
    type = tk.StringVar()
    type = ttk.Combobox(sub_win, values=[1, 2, 3], state="readonly")
    type.grid(row=3, column=1)
    type.current(0)

    def insert():
        try:
            c_id = course_id.get().strip()
            c_name = course_name.get().strip()
            t = type.get()
            insert_course(c_id, c_name, 0, t)
            messagebox.showinfo("成功", "插入课程成功！")
        except Exception as e:
            messagebox.showerror("错误", f"插入失败：{e}")
            sub_win.destroy()

    submit_btn = ttk.Button(sub_win, text="提交", command=insert)
    submit_btn.grid(row=6, column=0, columnspan=2, pady=20)

    sub_win.mainloop()

def del_course():
    sub_win = tk.Toplevel()
    sub_win.title("课程删除")
    sub_win.geometry("300x150")

    ttk.Label(sub_win, text="课程号：").grid(row=0, column=0, padx=10, pady=10, sticky="e")
    course_id = ttk.Entry(sub_win)
    course_id.grid(row=0, column=1)

    def delete():
        try:
            c_id = course_id.get().strip()
            delete_teacher_teach_course(current_teacher_id, c_id)
            messagebox.showinfo("成功", "删除成功！")
            if not hour_correct(c_id):
                messagebox.showinfo("提示", "课程总学时已更新")
        except Exception as e:
            messagebox.showerror("错误", f"删除失败：{e}")
            sub_win.destroy()

    submit_btn = ttk.Button(sub_win, text="提交", command=delete)
    submit_btn.grid(row=1, column=0, columnspan=2, pady=20)

    sub_win.mainloop()

def renewcourse():
    sub_win = tk.Toplevel()
    sub_win.title("课程修改")
    sub_win.geometry("300x500")
    ttk.Label(sub_win, text=f"请输入指定信息：", font=("Arial", 16))

    ttk.Label(sub_win, text="课程号：").grid(row=0, column=0, padx=10, pady=10, sticky="e")
    course_id = ttk.Entry(sub_win)
    course_id.grid(row=0, column=1)

    ttk.Label(sub_win, text="课程名称：").grid(row=1, column=0, padx=10, pady=10, sticky="e")
    course_name = ttk.Entry(sub_win)
    course_name.grid(row=1, column=1)

    ttk.Label(sub_win, text="学时数：").grid(row=2, column=0, padx=10, pady=10, sticky="e")
    hour = ttk.Entry(sub_win)
    hour.grid(row=2, column=1)

    ttk.Label(sub_win, text="课程性质：").grid(row=3, column=0, padx=10, pady=10, sticky="e")
    type = tk.StringVar()
    type = ttk.Combobox(sub_win, values=[1, 2, 3], state="readonly")
    type.grid(row=3, column=1)
    type.current(0)

    def renew():
        try:
            c_id = course_id.get().strip()
            c_name = course_name.get().strip()
            h = hour.get().strip()
            t = type.get()
            renew_course(c_id, c_name, h, t)
            messagebox.showinfo("成功", "更新课程成功！")
            if not hour_correct(c_id):
                messagebox.showinfo("提示", "课程总学时已更新")
        except Exception as e:
            messagebox.showerror("错误", f"更新失败：{e}")
            sub_win.destroy()

    submit_btn = ttk.Button(sub_win, text="提交", command=renew)
    submit_btn.grid(row=9, column=0, columnspan=2, pady=20)

    sub_win.mainloop()

def searchcourse():
    sub_win = tk.Toplevel()
    sub_win.title("课程查询")
    sub_win.geometry("500x500")
    ttk.Label(sub_win, text=f"请输入指定信息：", font=("Arial", 16))

    ttk.Label(sub_win, text="课程号：").grid(row=0, column=0, padx=10, pady=10, sticky="e")
    course_id = ttk.Entry(sub_win)
    course_id.grid(row=0, column=1)

    ttk.Label(sub_win, text="课程名称：").grid(row=1, column=0, padx=10, pady=10, sticky="e")
    course_name = ttk.Entry(sub_win)
    course_name.grid(row=1, column=1)

    ttk.Label(sub_win, text="学时数：").grid(row=2, column=0, padx=10, pady=10, sticky="e")
    hour = ttk.Entry(sub_win)
    hour.grid(row=2, column=1)

    ttk.Label(sub_win, text="课程性质：").grid(row=3, column=0, padx=10, pady=10, sticky="e")
    type = tk.StringVar()
    type = ttk.Combobox(sub_win, values=[1, 2, 3], state="readonly")
    type.grid(row=3, column=1)
    type.set('') 

    ttk.Label(sub_win, text="主讲教师工号：").grid(row=4, column=0, padx=10, pady=10, sticky="e")
    teacher_id = ttk.Entry(sub_win)
    teacher_id.grid(row=4, column=1)

    columns = ["course_id", "course_name", "hour", "course_type", "teacher_id"]
    tree = ttk.Treeview(sub_win, columns=columns, show="headings", height=10)
    tree.grid(row=6, column=0, columnspan=2, pady=10)

    headings = ["课程号", "课程名称", "学时数", "课程性质", "教师工号"]
    for col, heading in zip(columns, headings):
        tree.heading(col, text=heading)
        tree.column(col, width=100)

    def search():
        try:
            c_id = course_id.get().strip()
            c_name = course_name.get().strip()
            h = hour.get()
            t = type.get()
            t_id = teacher_id.get()
            results = search_course(c_id, c_name, h, t, t_id)

            for row in tree.get_children():
                tree.delete(row)
            for row in results:
                values = [row["course_id"], row["course_name"], row["hour"],
                        row["course_type"], row["teacher_id"]]
                tree.insert("", tk.END, values=values)
                # tree.insert("", tk.END, values=row)

        except Exception as e:
            messagebox.showerror("错误", f"查询失败：{e}")
            sub_win.destroy()

    submit_btn = ttk.Button(sub_win, text="提交", command=search)
    submit_btn.grid(row=5, column=0, columnspan=2, pady=20)

    sub_win.mainloop()

def search():
    sub_win = tk.Toplevel()
    sub_win.title("查询统计")
    sub_win.geometry("500x400")
    ttk.Label(sub_win, text=f"请输入想执行的查询：", font=("Arial", 16))

    ttk.Label(sub_win, text="年份下限：").grid(row=1, column=0, padx=10, pady=10, sticky="e")
    from_year = ttk.Entry(sub_win)
    from_year.grid(row=1, column=1)

    ttk.Label(sub_win, text="年份上限：").grid(row=2, column=0, padx=10, pady=10, sticky="e")
    to_year = ttk.Entry(sub_win)
    to_year.grid(row=2, column=1)

    def commit():
        from_year_val = int(from_year.get())
        to_year_val = int(to_year.get())

        try:
            result = get_teacher_statistics(current_teacher_id, from_year_val, to_year_val)

            if not result:
                messagebox.showinfo("无结果", "未查询到该教师在所选年份范围内的记录。")
                return

            detail_win = tk.Toplevel()
            detail_win.title("查询结果详情")
            detail_win.geometry("700x600")

            text = tk.Text(detail_win)
            text.pack(expand=True, fill='both')

            content = f"教师工号：{current_teacher_id}\n年份范围：{from_year_val} - {to_year_val}\n\n"

            content += f"发表论文：\n"
            for p in result['papers']:
                content += f"   论文序号：{p['paper_id']}，论文名称：{p['paper_name']}，发表源：{p['source']}，年份：{p['post_year']}，类型：{p['paper_type']}，级别：{p['level']}，排名：{p['teacher_paper_rank']}，是否通讯作者：{p['is_cor']}\n"

            content += f"主持项目：\n"
            for pr in result['projects']:
                content += f"   项目号：{pr['project_id']}，项目名称：{pr['project_name']}，项目来源：{pr['project_source']}，项目类型：{pr['project_type']}，起止年份：{pr['start_year']}-{pr['end_year']}，排名：{pr['teacher_project_rank']}，承担经费：{pr['teacher_hold_projectmoney']}\n"

            content += f"主讲课程：\n"
            for c in result['courses']:
                content += f"   课程号：{c['course_id']}，课程名：{c['course_name']}，课程类型：{c['course_type']}，年份：{c['year']}，学期：{c['term']}，承担学时：{c['own_hour']}\n"

            text.insert('1.0', content)
            text.config(state='disabled')

            def export_to_docx():
                file_path = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word 文档", "*.docx")]
                )
                doc = Document()
                style = doc.styles['Normal']
                font = style.font
                font.name = '宋体'
                font.size = Pt(12)
                style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                doc.add_heading('教师信息统计表', level=1)
                for line in content.strip().split('\n'):
                    doc.add_paragraph(line)
                doc.save(file_path)
                messagebox.showinfo("导出成功", f"文件已保存至：\n{file_path}")

            export_btn = ttk.Button(detail_win, text="导出为 Word 文档", command=export_to_docx)
            export_btn.pack(pady=10)
        
        except Exception as e:
                    messagebox.showerror("错误", f"查询失败：{e}")
                    sub_win.destroy()

    submit_btn = ttk.Button(sub_win, text="提交", command=commit)
    submit_btn.grid(row=3, column=0, columnspan=2, pady=20)

def view_table():
    sub_win = tk.Toplevel()
    sub_win.title("查看数据表内容")
    sub_win.geometry("600x400")

    ttk.Label(sub_win, text="请选择要查看的表：", font=("Arial", 12)).pack(pady=10)

    table_names = [
        "teachers", "papers", "projects", "courses",
        "teacher_post_paper", "teacher_hold_project", "teacher_teach_course"
    ]
    table_var = tk.StringVar()
    table_combo = ttk.Combobox(sub_win, textvariable=table_var, values=table_names, state="readonly")
    table_combo.pack(pady=5)
    table_combo.current(0)

    result_text = tk.Text(sub_win, height=20, width=70)
    result_text.pack(pady=10)

    def show_table():
        table = table_var.get()
        try:
            rows = fetch_all_from_table(table)
            result_text.delete(1.0, tk.END)

            if rows:
                col_names = rows[0].keys()
                result_text.insert(tk.END, " | ".join(col_names) + "\n")
                result_text.insert(tk.END, "-" * 60 + "\n")
                for row in rows:
                    row_str = " | ".join(str(value) for value in row.values())
                    result_text.insert(tk.END, row_str + "\n")
            else:
                result_text.insert(tk.END, f"{table} 表为空。")

        except Exception as e:
            result_text.delete(1.0, tk.END)
            result_text.insert(tk.END, f"查询失败：{e}")

    ttk.Button(sub_win, text="显示表内容", command=show_table).pack(pady=5)

## 主界面部分
def mainpage():
    root2 = tk.Toplevel()
    root2.title("教师教学科研登记系统")
    root2.geometry("800x600")

    ttk.Label(root2, text="请选择要进入的模块：", font=("Arial", 14)).pack(pady=20)

    ttk.Button(root2, text="发表论文情况", width=20, command=paper).pack(pady=5)
    ttk.Button(root2, text="承担项目情况", width=20, command=project).pack(pady=5)
    ttk.Button(root2, text="主讲课程情况", width=20, command=course).pack(pady=5)
    ttk.Button(root2, text="查询统计", width=20, command=search).pack(pady=5)
    ttk.Button(root2, text="查看数据表", width=20, command=view_table).pack(pady=5)

    root2.mainloop()

def login_window():
    login_win = tk.Toplevel()
    login_win.title("教师登录")
    login_win.geometry("300x200")

    ttk.Label(login_win, text="教师工号：").pack(pady=5)
    id_entry = ttk.Entry(login_win)
    id_entry.pack()

    ttk.Label(login_win, text="密码：").pack(pady=5)
    pwd_entry = ttk.Entry(login_win, show="*")
    pwd_entry.pack()

    def do_login():
        global current_teacher_id
        teacher_id = id_entry.get()
        password = pwd_entry.get()
        if check_teacher_password(teacher_id, password):
            current_teacher_id = teacher_id
            messagebox.showinfo("登录成功", f"欢迎，{teacher_id}")
            mainpage()
        else:
            messagebox.showerror("登录失败", "工号或密码错误")
            login_win.destroy()

    ttk.Button(login_win, text="登录", command=do_login).pack(pady=10)

def main():
    root = tk.Tk()
    root.title("教师教学科研登记系统")
    root.geometry("800x600")

    ttk.Label(root, text="请先登入：", font=("Arial", 14)).pack(pady=20)

    ttk.Button(root, text="登录", width=20, command=login_window).pack(pady=5)

    root.mainloop()

main()
